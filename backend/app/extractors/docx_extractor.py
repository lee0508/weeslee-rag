"""
DOCX Extractor
"""
import os
from typing import Dict, Any, List
from pathlib import Path
from docx import Document

from app.extractors.base import BaseExtractor, ExtractionResult


class DocxExtractor(BaseExtractor):
    """Extracts text from DOCX files"""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx"]

    async def extract(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from DOCX file

        Args:
            file_path: Path to DOCX file

        Returns:
            Extraction result dictionary
        """
        if not os.path.exists(file_path):
            return ExtractionResult(
                success=False,
                error=f"File not found: {file_path}"
            ).to_dict()

        try:
            doc = Document(file_path)
            content_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                if table_text:
                    content_parts.append("\n".join(table_text))

            content = "\n\n".join(content_parts)

            metadata = {
                "source": file_path,
                "filename": Path(file_path).name,
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            }

            # Extract document properties if available
            try:
                core_props = doc.core_properties
                if core_props.title:
                    metadata["title"] = core_props.title
                if core_props.author:
                    metadata["author"] = core_props.author
                if core_props.created:
                    metadata["created"] = str(core_props.created)
            except Exception:
                pass

            return ExtractionResult(
                success=True,
                content=content,
                metadata=metadata,
                method="python-docx"
            ).to_dict()

        except Exception as e:
            return ExtractionResult(
                success=False,
                error=str(e),
                method="python-docx"
            ).to_dict()


# Singleton instance
docx_extractor = DocxExtractor()
